from google.adk.agents import LlmAgent
from google.adk.tools import google_search, VertexAiSearchTool
from google import genai
from google.genai.types import GenerateVideosConfig
import os
import re
import sys
import time
from dotenv import load_dotenv
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from pathlib import Path
import vertexai
from vertexai.preview.vision_models import ImageGenerationModel
from sahayak_agent.utils import (
    store_file,
    clean_text,
    add_text_with_newlines,
    QuizData,
    WorksheetData,
    LessonPlanData,
    StructuredEducationalContent,
    ContentType
)

load_dotenv()

BUCKET_ID = 'sahayak-user-documents'

vertexai.init(project=os.getenv('GOOGLE_CLOUD_PROJECT'), location=os.getenv('GOOGLE_CLOUD_LOCATION'))

search_agent = LlmAgent(
    name="search_agent",
    model="gemini-2.0-flash-001",
    description="""You are 'search_agent'.
    You help with searching and retrieving information from various sources including:
    1. NCERT textbooks and educational content via Vertex AI Search
    2. General web search for additional context and information
    """,
    instruction=f"""
    You have access to two search tools:
    1. vertex_ai_search - for searching NCERT textbooks and educational content
    2. google_search - for general web searches
    
    Your job is to:
    1. Search for relevant educational content based on user requirements
    2. Gather comprehensive information about the requested topic
    3. Provide detailed, well-researched content that can be used for creating educational materials
    
    Always provide detailed, accurate, and helpful responses based on the search results.
    Format your responses in a clear and educational manner.
    Focus on gathering ALL necessary information for the educational content being requested.
    
    IMPORTANT:
    - Provide comprehensive information that covers all aspects of the topic
    - Include examples, explanations, and educational context
    - Clean up any malformed content before returning
    - Simple text formatting is allowed, but no HTML or Markdown
    - If a search tool fails, try the other one or provide content from your knowledge
    - Keep responses concise but comprehensive to avoid timeout issues
    """,
    tools=[
        google_search,
        VertexAiSearchTool(
            data_store_id=f"projects/{os.getenv('GOOGLE_CLOUD_PROJECT')}/locations/global/collections/default_collection/dataStores/{os.getenv('GOOGLE_CLOUD_DATASTORE_ID')}"
        )
    ]
)

from google import genai
from google.genai.types import GenerateContentConfig

def structure_data_tool(request: str) -> dict:
    """Create structured educational content using direct Gemini API call with structured output.
    
    Args:
        request (str): The request to format educational content into structured format. Must say what type of content is being requested (quiz, worksheet, lesson plan).
    
    Returns:
        dict: The structured educational content.
    """
    client = genai.Client()
    
    model_id = "gemini-2.0-flash-001"
    
    instruction = """
    You will receive educational content and a request to format it into one of these structures:
    1. Quiz - Multiple choice questions with correct answers
    2. Lesson Plan - Structured lessons with content and duration
    3. Worksheet - Questions with detailed answers
    
    Your job is to:
    1. Analyze the provided educational content
    2. Create the requested structured output based on the content
    3. Ensure all content is educationally sound and age-appropriate
    4. Follow NEP guidelines for lesson plans
    5. Do not change the content, only format it.
    
    Guidelines:
    - For quizzes: Provide at least 10 questions by default, each with at least 2 options (default 4)
    - For worksheets: Provide at least 5 questions by default with descriptive answers
    - For lesson plans: Follow NEP guidelines and provide detailed, structured lessons
    - Ensure all content is clear, accurate, and educationally valuable
    - Use simple language appropriate for the specified grade level
    
    IMPORTANT: Always set the content_type field to indicate what type of content you're creating:
    - Set content_type to "quiz" and fill quiz_data for quizzes
    - Set content_type to "worksheet" and fill worksheet_data for worksheets  
    - Set content_type to "lesson_plan" and fill lesson_plan_data for lesson plans
    
    The content_type must be one of: "quiz", "worksheet", or "lesson_plan"
    """
    
    response = client.models.generate_content(
        model=model_id,
        contents=f"""
        SYSTEM: {instruction}
        
        user: {request}
        """,
        config=GenerateContentConfig(
            response_schema=StructuredEducationalContent,
            response_mime_type="application/json"
        )
    )
    return response.to_json_dict()

def quiz_creator(quiz_data: dict, user_id: str) -> str:
    """Create a DOCX file from quiz data.
    
    Args:
        quiz_data (dict): The quiz data as a dictionary.
        user_id (str): The ID of the user for whom the quiz is being created.
    
    Returns:
        str: The file path of the created DOCX file.
    """
    try:
        if isinstance(quiz_data, dict):
            quiz_data = QuizData(**quiz_data)
        elif not isinstance(quiz_data, QuizData):
            return f"Invalid quiz_data type: {type(quiz_data)}"
            
        doc = Document()
        title = doc.add_heading(clean_text(quiz_data.quiz_title), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        subtitle = doc.add_heading(clean_text(quiz_data.quiz_subtitle), level=1)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        doc.add_paragraph("Instructions: Choose the correct answer(s) for each question.")
        doc.add_paragraph()

        for i, question in enumerate(quiz_data.questions, 1):
            question_para = doc.add_paragraph(f"Q{i}. ")
            add_text_with_newlines(question_para, clean_text(question.question))
            question_para.style = 'Heading 3'
            
            for j, option in enumerate(question.options, 1):
                option_para = doc.add_paragraph(f"   {chr(64+j)}. {clean_text(option)}")
            
            doc.add_paragraph()
        
        doc.add_page_break()
        doc.add_heading("Answer Key", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for i, question in enumerate(quiz_data.questions, 1):
            correct_letters = [chr(65 + idx) for idx in question.correct_answers]
            answer_para = doc.add_paragraph(f"Q{i}. {', '.join(correct_letters)}")
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        clean_title = re.sub(r'[^\w\s-]', '', quiz_data.quiz_title).strip()[:50]
        clean_title = re.sub(r'[-\s]+', '_', clean_title)
        filename = f"quiz_{clean_title}_{timestamp}.docx"
        filepath = Path("temp") / filename
        filepath.parent.mkdir(parents=True, exist_ok=True)
        
        doc.save(str(filepath))
        
        if not filepath.exists():
            return f"Failed to create file at {filepath}"

        stored_path = store_file(str(filepath), user_id)
        return str(stored_path)
        
    except Exception as e:
        print(f"Error in quiz_creator: {e}")
        print(f"Data type received: {type(quiz_data)}")
        if hasattr(quiz_data, '__dict__'):
            print(f"Data content: {quiz_data.__dict__}")
        else:
            print(f"Data content: {quiz_data}")
        return f"I apologize, but there was an error generating your quiz. Please try again later. Error: {e}"

def worksheet_creator(worksheet_data: dict, user_id: str) -> str:
    """Create a DOCX file from worksheet data.
    
    Args:
        worksheet_data (dict): The worksheet data as a dictionary.
        user_id (str): The ID of the user for whom the worksheet is being created.
    
    Returns:
        str: The file path of the created DOCX file.
    """
    try:
        # Convert dictionary to Pydantic model
        if isinstance(worksheet_data, dict):
            worksheet_data = WorksheetData(**worksheet_data)
        elif not isinstance(worksheet_data, WorksheetData):
            return f"Invalid worksheet_data type: {type(worksheet_data)}"
            
        doc = Document()  
        title = doc.add_heading(clean_text(worksheet_data.worksheet_title), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        subtitle = doc.add_heading(clean_text(worksheet_data.worksheet_subtitle), level=1)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        doc.add_paragraph("Instructions: Answer the following questions in detail.")
        doc.add_paragraph()

        for i, question in enumerate(worksheet_data.questions, 1):
            question_para = doc.add_paragraph(f"Q{i}. ")
            add_text_with_newlines(question_para, clean_text(question.question))
            question_para.style = 'Heading 3'
            
            answer_para = doc.add_paragraph("Answer: ")
            add_text_with_newlines(answer_para, clean_text(question.answer))
            doc.add_paragraph()
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        clean_title = re.sub(r'[^\w\s-]', '', worksheet_data.worksheet_title).strip()[:50]
        clean_title = re.sub(r'[-\s]+', '_', clean_title)
        filename = f"worksheet_{clean_title}_{timestamp}.docx"
        filepath = Path("temp") / filename
        filepath.parent.mkdir(parents=True, exist_ok=True)
        
        doc.save(str(filepath))
        
        if not filepath.exists():
            return f"Failed to create file at {filepath}"

        stored_path = store_file(str(filepath), user_id)
        return str(stored_path)
        
    except Exception as e:
        print(f"Error in worksheet_creator: {e}")
        print(f"Data type received: {type(worksheet_data)}")
        if hasattr(worksheet_data, '__dict__'):
            print(f"Data content: {worksheet_data.__dict__}")
        else:
            print(f"Data content: {worksheet_data}")
        return f"I apologize, but there was an error generating your worksheet. Please try again later. Error: {e}"

def lesson_plan_creator(lesson_plan_data: dict, user_id: str) -> str:
    """Create a DOCX file for a lesson plan.

    Args:
        lesson_plan_data (dict): The lesson plan data as a dictionary.
        user_id (str): The ID of the user for whom the lesson plan is being created.

    Returns:
        str: The file path of the created DOCX file.
    """
    try:
        # Convert dictionary to Pydantic model
        if isinstance(lesson_plan_data, dict):
            lesson_plan_data = LessonPlanData(**lesson_plan_data)
        elif not isinstance(lesson_plan_data, LessonPlanData):
            return f"Invalid lesson_plan_data type: {type(lesson_plan_data)}"
        
        doc = Document()
        title = doc.add_heading(clean_text(lesson_plan_data.lesson_plan_title), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        subtitle = doc.add_heading(clean_text(lesson_plan_data.lesson_plan_description), level=1)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        doc.add_paragraph("Instructions: Follow the lesson plan as outlined below.")
        doc.add_paragraph()
        
        for i, lesson in enumerate(lesson_plan_data.lessons, 1):
            lesson_para = doc.add_paragraph(f"{lesson.lesson_title}")
            lesson_para.style = 'Heading 3'

            if lesson.lesson_details:
                table = doc.add_table(rows=1, cols=2)
                table.style = 'Light List Accent 1'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Content'
                hdr_cells[1].text = 'Duration'
                
                for detail in lesson.lesson_details:
                    row_cells = table.add_row().cells
                    content_lines = clean_text(detail.lesson_content).split('\n')
                    row_cells[0].text = '\n'.join(content_lines)
                    row_cells[1].text = clean_text(detail.duration)

                for row in table.rows:
                    for cell in row.cells:
                        cell.width = Inches(3)

                doc.add_paragraph()
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        clean_title = re.sub(r'[^\w\s-]', '', lesson_plan_data.lesson_plan_title).strip()[:50]
        clean_title = re.sub(r'[-\s]+', '_', clean_title)
        filename = f"lesson_plan_{clean_title}_{timestamp}.docx"
        filepath = Path("temp") / filename
        filepath.parent.mkdir(parents=True, exist_ok=True)
        
        doc.save(str(filepath))
        
        if not filepath.exists():
            return f"Failed to create file at {filepath}"
        
        stored_path = store_file(str(filepath), user_id)
        return str(stored_path)
        
    except Exception as e:
        return f"I apologize, but there was an error generating your lesson plan. Please try again later. Error: {e}"

def image_generator(prompt: str, user_id: str) -> str:
    """Generate an image from a text prompt using Vertex AI.

    Args:
        prompt (str): The text prompt to generate an image for.
        user_id (str): The ID of the user for whom the image is being generated.

    Returns:
        str: The file path of the generated image.
    """
    import base64
    
    try:
        model = ImageGenerationModel.from_pretrained("imagen-4.0-fast-generate-preview-06-06")
        images = model.generate_images(
            prompt=prompt + "\n Make sure image is simple line drawings or flow charts, not fancy illustrations.",
            number_of_images=1
        )

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_prompt = re.sub(r'[^\w\s-]', '', prompt).strip()[:50]
        safe_prompt = re.sub(r'[-\s]+', '_', safe_prompt)
        filename = f"generated_image_{safe_prompt}_{timestamp}.png"
        filepath = Path("temp") / filename
        filepath.parent.mkdir(exist_ok=True)
        
        image_data = base64.b64decode(images[0]._as_base64_string())
        with open(filepath, 'wb') as f:
            f.write(image_data)
        
        stored_path = store_file(filepath, user_id)
        return str(stored_path)

    except Exception as e:
        return f"Error generating image: {e}"

def video_generator(prompt: str, user_id: str) -> str:
    """Generate a video from a text prompt using Vertex AI.

    Args:
        prompt (str): The text prompt to generate a video for.
        user_id (str): The ID of the user for whom the video is being generated.

    Returns:
        str: The file path of the generated video.
    """
    try:
        global BUCKET_ID
        client = genai.Client()
        output_gcs_uri = f"gs://{BUCKET_ID}/videos/{user_id}/{prompt[:50].replace(' ', '_')}/"

        operation = client.models.generate_videos(
            model="veo-3.0-generate-preview",
            prompt=prompt,
            config=GenerateVideosConfig(
                aspect_ratio="16:9",
                output_gcs_uri=output_gcs_uri,
            ),
        )

        while not operation.done:
            time.sleep(15)
            operation = client.operations.get(operation)

        if operation.response:
            return operation.result.generated_videos[0].video.uri

    except Exception as e:
        return f"Error generating video: {e}"